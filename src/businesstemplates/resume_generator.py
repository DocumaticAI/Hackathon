"""
Copyright (c) 2022 That GenZ Gamer <thatgenzgamer@gmail.com>

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the “Software”), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS
OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from .utils import change_text, find_paragraph, insert_paragraph_after
from pathlib import Path


def main():
	# Prepare a nice (custom) ASCII art to show before the start of the inputs.
	print("--------------------------------")
	print("|    Python Resume Generator    |")
	print("|    -----------------------    |")
	print("| Now, you can design your own  |")
	print("|      impressive resume.       |")
	print("--------------------------------\n")

	template_directory = Path(__file__).parent / "templates"
	resume_template = template_directory / "Resume Template.docx"

	# Reading the document containing the resume template.
	document = Document(resume_template)

	# The following set of inputs is to get enough information to prepare the resume.

	# Getting the user's name
	name = input("What is your full name?: ")

	# Changing the text of the user's name and setting the font size using an helper function.
	if len(name) != 0:
		change_text(document.paragraphs[0], name, 18, True)
	else:
		print("You haven't mentioned a name!")
		exit()

	# Getting the information table from the word document.
	information_table = document.tables[0]

	# Asking for the user's age
	age = input("What is your age?: ")

	# Changing the text of the user's age and setting the font size using an helper function.
	if len(age) != 0:
		age_cell = information_table.columns[0].cells[0].paragraphs[0]

		change_text(age_cell, f"Age: {age}", 12)
	else:
		print("You haven't mentioned an age!")
		exit()

	# Asking for the user's gender
	gender = input("What is your gender?: ")

	# Changing the text of the user's gender and setting the font size using an helper function.
	if len(gender) != 0:
		gender_cell = information_table.columns[0].cells[1].paragraphs[0]

		change_text(gender_cell, f"Gender: {gender}", 12)
	else:
		print("You haven't mentioned a gender!")
		exit()

	# Asking for the user's nationality
	nationality = input("Which country do you belong to?: ")

	# Changing the text of the user's nationality and setting the font size using an helper function.
	if len(nationality) != 0:
		nationality_cell = information_table.columns[0].cells[2].paragraphs[0]

		change_text(nationality_cell, f"Nationality: {nationality}", 12)
	else:
		print("You haven't mentioned a nationality!")
		exit()

	# Asking for the user's Email ID
	email = input("What is your Email ID?: ")

	# Changing the text of the user's Email ID and setting the font size using an helper function.
	if len(email) != 0:
		email_cell = information_table.columns[1].cells[0].paragraphs[0]

		change_text(email_cell, f"Email ID: {email}", 12)
	else:
		print("You haven't mentioned an Email ID!")
		exit()

	# Asking for the user's Mobile no
	mobile_number = input("What is your mobile number?: ")

	# Changing the text of the user's mobile no and setting the font size using an helper function.
	if len(mobile_number) != 0:
		mobile_number_cell = information_table.columns[1].cells[1].paragraphs[0]

		change_text(mobile_number_cell, f"Mobile No: {mobile_number}", 12)
	else:
		print("You haven't entered a mobile number")
		exit()

	# Asking for the user's city
	city = input("Which city are you from?: ")

	# Changing the text of the user's mobile no and setting the font size using an helper function.
	if len(city) != 0:
		city_cell = information_table.columns[1].cells[2].paragraphs[0]

		change_text(city_cell, f"City: {city}", 12)
	else:
		print("You haven't mentioned a city!")
		exit()

	# Asking for the user's introduction
	objective = input("Write your objective to be mentioned in the resume: ")

	# Changing the text of the user's introduction and setting the font size using an helper function.
	if len(objective) != 0:
		change_text(document.paragraphs[3], objective, 12)
	else:
		print("You haven't mentioned your objective!")
		exit()

	# Getting the course details table
	course_details = document.tables[1]

	# Asking for the highest qualification
	highest_qualification = input(
		"What is your highest qualification? [Employee/Post Graduation/Graduation/Class 12]: "
	).lower()

	qualifications = [
		"employee", 
		"post graduation", 
		"graduation", 
		"class 12",
	]

	# Exitting if the qualification selected does not exist in options.
	if highest_qualification not in qualifications:
		print("Pick an option from the given options!")
		exit()
	# Exitting if the user hasn't entered anything.
	elif len(highest_qualification) == 0:
		print("You haven't mentioned your highest qualification!")
	else:
		# Asking for post graduation, graduation, 12th and 10th details
		# if the user is an employee.
		if highest_qualification == "employee":
			employee_pg_institute = input(
				"Which institute did you pursue your post graduation in?: "
			)

			change_text(
				course_details.columns[1].cells[1].paragraphs[0],
				employee_pg_institute,
				12,
			)

			employee_pg_year = input(
				"Which year did you complete your post graduation?: "
			)

			change_text(
				course_details.columns[2].cells[1].paragraphs[0],
				employee_pg_year,
				12,
			)

			employee_pg_percentage = input(
				"What was your post graduation percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[1].paragraphs[0],
				employee_pg_percentage,
				12,
			)

			employee_grad_institute = input(
				"Which institute did you pursue your graduation in?: "
			)

			change_text(
				course_details.columns[1].cells[2].paragraphs[0],
				employee_grad_institute,
				12,
			)

			employee_grad_year = input(
				"Which year did you complete your graduation?: "
			)

			change_text(
				course_details.columns[2].cells[2].paragraphs[0],
				employee_grad_year,
				12,
			)

			employee_grad_percentage = input(
				"What was your graduation percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[2].paragraphs[0],
				employee_grad_percentage,
				12,
			)

			employee_12th_institute = input(
				"Which institute did you pursue your 12th grade in?: "
			)

			change_text(
				course_details.columns[1].cells[3].paragraphs[0],
				employee_12th_institute,
				12,
			)

			employee_12th_year = input(
				"Which year did you complete your 12th grade?: "
			)

			change_text(
				course_details.columns[2].cells[3].paragraphs[0],
				employee_12th_year,
				12,
			)

			employee_12th_percentage = input(
				"What was your 12th grade percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[3].paragraphs[0],
				employee_12th_percentage,
				12,
			)

			employee_10th_institute = input(
				"Which institute did you pursue your 10th grade in?: "
			)

			change_text(
				course_details.columns[1].cells[4].paragraphs[0],
				employee_10th_institute,
				12,
			)

			employee_10th_year = input(
				"Which year did you complete your 10th grade?: "
			)

			change_text(
				course_details.columns[2].cells[4].paragraphs[0],
				employee_10th_year,
				12,
			)

			employee_10th_percentage = input(
				"What was your 10th grade percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[4].paragraphs[0],
				employee_10th_percentage,
				12,
			)

		# Asking for pg, ug, 12th and 10th details if user is an post graduate.
		elif highest_qualification == "post graduation":
			postgraduate_pg_institute = input(
				"Which institute did you pursue your post graduation in?: "
			)

			change_text(
				course_details.columns[1].cells[1].paragraphs[0],
				postgraduate_pg_institute,
				12,
			)

			postgraduate_pg_year = input(
				"Which year did you complete your post graduation?: "
			)

			change_text(
				course_details.columns[2].cells[1].paragraphs[0],
				postgraduate_pg_year,
				12,
			)

			postgraduate_pg_percentage = input(
				"What was your post graduation percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[1].paragraphs[0],
				postgraduate_pg_percentage,
				12,
			)

			postgraduate_grad_institute = input(
				"Which institute did you pursue your graduation in?: "
			)

			change_text(
				course_details.columns[1].cells[2].paragraphs[0],
				postgraduate_grad_institute,
				12,
			)

			postgraduate_grad_year = input(
				"Which year did you complete your graduation?: "
			)

			change_text(
				course_details.columns[2].cells[2].paragraphs[0],
				postgraduate_grad_year,
				12,
			)

			postgraduate_grad_percentage = input(
				"What was your graduation percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[2].paragraphs[0],
				postgraduate_grad_percentage,
				12,
			)

			postgraduate_12th_institute = input(
				"Which institute did you pursue your 12th grade in?: "
			)

			change_text(
				course_details.columns[1].cells[3].paragraphs[0],
				postgraduate_12th_institute,
				12,
			)

			postgraduate_12th_year = input(
				"Which year did you complete your 12th grade?: "
			)

			change_text(
				course_details.columns[2].cells[3].paragraphs[0],
				postgraduate_12th_year,
				12,
			)

			postgraduate_12th_percentage = input(
				"What was your 12th grade percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[3].paragraphs[0],
				postgraduate_12th_percentage,
				12,
			)

			postgraduate_10th_institute = input(
				"Which institute did you pursue your 10th grade in?: "
			)

			change_text(
				course_details.columns[1].cells[4].paragraphs[0],
				postgraduate_10th_institute,
				12,
			)

			postgraduate_10th_year = input(
				"Which year did you complete your 10th grade?: "
			)

			change_text(
				course_details.columns[2].cells[4].paragraphs[0],
				postgraduate_10th_year,
				12,
			)

			postgraduate_10th_percentage = input(
				"What was your 10th grade percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[4].paragraphs[0],
				postgraduate_10th_percentage,
				12,
			)

		# Asking for ug, 12th and 10th details if user is an graduate.
		elif highest_qualification == "graduation":
			change_text(course_details.columns[1].cells[1].paragraphs[0], "NA", 12)
			change_text(course_details.columns[2].cells[1].paragraphs[0], "NA", 12)
			change_text(course_details.columns[3].cells[1].paragraphs[0], "NA", 12)

			graduate_grad_institute = input(
				"Which institute did you pursue your graduation in?: "
			)

			change_text(
				course_details.columns[1].cells[2].paragraphs[0],
				graduate_grad_institute,
				12,
			)

			graduate_grad_year = input(
				"Which year did you complete your graduation?: "
			)

			change_text(
				course_details.columns[2].cells[2].paragraphs[0],
				graduate_grad_year,
				12,
			)

			graduate_grad_percentage = input(
				"What was your graduation percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[2].paragraphs[0],
				graduate_grad_percentage,
				12,
			)

			graduate_12th_institute = input(
				"Which institute did you pursue your 12th grade in?: "
			)

			change_text(
				course_details.columns[1].cells[3].paragraphs[0],
				graduate_12th_institute,
				12,
			)

			graduate_12th_year = input(
				"Which year did you complete your 12th grade?: "
			)

			change_text(
				course_details.columns[2].cells[3].paragraphs[0],
				graduate_12th_year,
				12,
			)

			graduate_12th_percentage = input(
				"What was your 12th grade percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[3].paragraphs[0],
				graduate_12th_percentage,
				12,
			)

			graduate_10th_institute = input(
				"Which institute did you pursue your 10th grade in?: "
			)

			change_text(
				course_details.columns[1].cells[4].paragraphs[0],
				graduate_10th_institute,
				12,
			)

			graduate_10th_year = input(
				"Which year did you complete your 10th grade?: "
			)

			change_text(
				course_details.columns[2].cells[4].paragraphs[0],
				graduate_10th_year,
				12,
			)

			graduate_10th_percentage = input(
				"What was your 10th grade percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[4].paragraphs[0],
				graduate_10th_percentage,
				12,
			)

		# Asking for 12th and 10th details if user is a 12th grader.
		elif highest_qualification == "class 12":
			change_text(course_details.columns[1].cells[1].paragraphs[0], "NA", 12)
			change_text(course_details.columns[2].cells[1].paragraphs[0], "NA", 12)
			change_text(course_details.columns[3].cells[1].paragraphs[0], "NA", 12)

			change_text(course_details.columns[1].cells[2].paragraphs[0], "NA", 12)
			change_text(course_details.columns[2].cells[2].paragraphs[0], "NA", 12)
			change_text(course_details.columns[3].cells[2].paragraphs[0], "NA", 12)

			class12_12th_institute = input(
				"Which institute did you pursue your 12th grade in?: "
			)

			change_text(
				course_details.columns[1].cells[3].paragraphs[0],
				class12_12th_institute,
				12,
			)

			class12_12th_year = input(
				"Which year did you complete your 12th grade?: "
			)

			change_text(
				course_details.columns[2].cells[3].paragraphs[0],
				class12_12th_year,
				12,
			)

			class12_12th_percentage = input(
				"What was your 12th grade percentage obtained?: "
			)
			change_text(
				course_details.columns[3].cells[3].paragraphs[0],
				class12_12th_percentage,
				12,
			)

			class12_10th_institute = input(
				"Which institute did you pursue your 10th grade in?: "
			)

			change_text(
				course_details.columns[1].cells[4].paragraphs[0],
				class12_10th_institute,
				12,
			)

			class12_10th_year = input(
				"Which year did you complete your 10th grade?: "
			)

			change_text(
				course_details.columns[2].cells[4].paragraphs[0],
				class12_10th_year,
				12,
			)

			class12_10th_percentage = input(
				"What was your 10th grade percentage obtained?: "
			)

			change_text(
				course_details.columns[3].cells[4].paragraphs[0],
				class12_10th_percentage,
				12,
			)
		else:
			pass

	# Asking the user whether he/she has previous work experience
	has_work_experience = input(
		"Do you have any previous work experience? [Yes/No]: "
	).lower()

	# Checking if he/she has work experience or not
	if has_work_experience == "yes":
		# Asking for the number of work experiences.
		number_of_work_experiences = input(
			"How many companies have you worked for?: "
		)

		# Checking if the user has provided a valid number.
		if number_of_work_experiences.isnumeric():
			number_of_work_experiences = int(number_of_work_experiences)
		else:
			print(
				"You haven't entered a valid number. Check if the number you have entered doesn't contain negative values or decimals and try again!"
			)
			exit()

		# Getting the first paragraph using an helper tool.
		work_experience_previous_paragraph = find_paragraph(document, "WORK EXPERIENCE")

		# Adding the paragraphs after one another containing the work experience information.
		for _ in range(0, number_of_work_experiences):
			work_experience = input(
				"Mention a suitable description for your previous work experience: "
			)

			work_experience_inserted_paragraph = insert_paragraph_after(
				document,
				work_experience_previous_paragraph,
				12,
				work_experience,
			)
			work_experience_previous_paragraph = work_experience_inserted_paragraph

	# Passing if the user doesn't have any work experience
	elif has_work_experience == "no":
		insert_paragraph_after(
			document,
			find_paragraph(document, "WORK EXPERIENCE"),
			12,
			"NA"
		)
	else:
		print("You haven't mentioned if you have previous work experience or not!")
		exit()

	# Asking the user whether he/she has previous internship experience.
	has_internship_experience = input(
		"Do you have any previous Internship experience? [Yes/No]: "
	).lower()

	# Checking if he/she has previous internship experience or not.
	if has_internship_experience == "yes":
		# Asking for the number of internship experience(s).
		number_of_internship_experiences = input(
			"How many internships have you attended?: "
		)

		# Checking if the user has provided a valid number.
		if number_of_internship_experiences.isnumeric():
			number_of_internship_experiences = int(
				number_of_internship_experiences
			)
		else:
			print(
				"You haven't entered a valid number. Check if the number you have entered doesn't contain negative values or decimals and try again!"
			)
			exit()

		# Getting the first paragraph using an helper tool.
		internship_experience_previous_paragraph = find_paragraph(
			document, "INTERNSHIP EXPERIENCE(S) (IF ANY)"
		)

		# Adding the paragraphs after one another containing the internship experience information.
		for _ in range(0, number_of_internship_experiences):
			internship_experience = input(
				"Mention a description for your previous work experience: "
			)

			internship_experience_inserted_paragraph = insert_paragraph_after(
				document,
				internship_experience_previous_paragraph,
				12,
				internship_experience,
			)
			internship_experience_previous_paragraph = (
				internship_experience_inserted_paragraph
			)

	# Passing if the user doesn't have any previous internship experience
	elif has_internship_experience == "no":
		insert_paragraph_after(
			document,
			find_paragraph(
				document, "INTERNSHIP EXPERIENCE(S) (IF ANY)"
			),
			12,
			"NA"
		)
	else:
		print("You haven't mentioned if you have an internship experience or not!")
		exit()

	# Asking the user whether he/she has certificates/courses.
	has_certificates_courses = input(
		"Do you have any previous Certificates/Courses to your name? [Yes/No]: "
	).lower()

	# Checking if he/she has any certificates/courses or not.
	if has_certificates_courses == "yes":
		# Splitting the certificates and courses with a ","
		certificates_courses = input(
			"Mention all the certificates/courses' name seperated by commas (,). Make sure to not include a space after each comma: "
		).split(",")

		# Getting the first paragraph using a helper tool.
		certificates_courses_previous_paragraph = find_paragraph(
			document, "CERTIFICATIONS/COURSES COMPLETED (IF ANY)"
		)

		# Adding the paragraphs after one another containing the certificate/course information.
		for certificate_course in certificates_courses:
			certificates_courses_inserted_paragraph = insert_paragraph_after(
				document=document,
				paragraph=certificates_courses_previous_paragraph,
				font_size=12,
				description=certificate_course,
			)

			certificates_courses_previous_paragraph = (
				certificates_courses_inserted_paragraph
			)

	# Passing if the user doesn't have any certificate/course
	elif has_certificates_courses == "no":
		insert_paragraph_after(
			document,
			find_paragraph(
				document, "CERTIFICATIONS/COURSES COMPLETED (IF ANY)"
			),
			12,
			"NA"
		)
	else:
		print("You haven't mentioned if you have a certificate/course or not!")
		exit()

	# Asking the user whether he/she has any hobby.
	has_skillsets = input("Do you have any skillsets? [Yes/No]: ").lower()

	# Checking if he/she has any skillsets or not.
	if has_skillsets == "yes":
		# Asking for the skillsets and seperating it by a ","
		skillsets = input(
			"Mention all the skillsets' name seperated by commas (,). Make sure to not include a space after each comma: "
		).split(",")

		# Getting the first paragraph.
		skillsets_previous_paragraph = find_paragraph(document, "SKILL SETS")

		# Adding the paragraphs after one another containing the skillset information.
		for skillset in skillsets:
			skillsets_inserted_paragraph = insert_paragraph_after(
				document=document,
				paragraph=skillsets_previous_paragraph,
				font_size=12,
				description=skillset,
			)

			skillsets_previous_paragraph = skillsets_inserted_paragraph

	# Passing if the user doesn't have any skillset
	elif has_skillsets == "no":
		insert_paragraph_after(
			document,
			find_paragraph(document, "SKILL SETS"),
			12,
			"NA"
		)
	else:
		print("You haven't mentioned if you have a skillset or not!")
		exit()

	# Asking the user whether he/she has any hobbies.
	has_hobbies = input("Do you have any hobbies? [Yes/No]: ").lower()

	# Checking if he/she has any hobby or not.
	if has_hobbies == "yes":
		# Asking for the hobbies
		hobbies = input(
			"Mention all the hobbies' name seperated by commas (,). Make sure to not include a space after each comma: "
		).split(",")

		# Getting the first paragraph.
		hobbies_previous_paragraph = find_paragraph(document, "HOBBIES AND INTERESTS")

		# Adding the paragraphs after one another containing the hobbies information.
		for hobby in hobbies:
			hobbies_inserted_paragraph = insert_paragraph_after(
				document=document,
				paragraph=hobbies_previous_paragraph,
				font_size=12,
				description=hobby,
			)

			hobbies_previous_paragraph = hobbies_inserted_paragraph

	# Passing if the user doesn't have any hobby
	elif has_hobbies == "no":
		insert_paragraph_after(
			document,
			find_paragraph(document, "HOBBIES AND INTERESTS"),
			12,
			"NA"
		)
	else:
		print("You haven't mentioned if you have a hobby or not!")
		exit()

	# Asking for the path to save the end document.
	path_to_save = input(
		"Where do you want to save the file? (Eg: C:/path/to/save/filename.docx) [Note: Relative paths are also allowed]: "
	)
	try:
		document.save(path_to_save)
	except:
		print("You have entered an incorrect path. Try again")
		exit()

# Running the script with the __name__ == "__main__" check.
if __name__ == "__main__":
	main()
